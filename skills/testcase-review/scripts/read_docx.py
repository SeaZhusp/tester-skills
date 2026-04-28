#!/usr/bin/env python3
"""
读取Word (.docx) 需求文档，提取文本内容
"""

import json
import sys
import argparse
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


def read_docx(file_path: str) -> dict:
    """
    读取Word文档

    Args:
        file_path: docx文件路径

    Returns:
        包含文档内容和结构的字典
    """
    doc = Document(file_path)

    content = []
    current_section = None
    sections = {}

    # 提取段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 判断是否为标题（根据样式）
        if para.style.name.startswith('Heading'):
            current_section = text
            if current_section not in sections:
                sections[current_section] = []
        elif current_section:
            sections[current_section].append(text)
        else:
            # 无标题的正文
            if '正文' not in sections:
                sections['正文'] = []
            sections[current_section].append(text)

        content.append({
            'type': 'heading' if para.style.name.startswith('Heading') else 'paragraph',
            'style': para.style.name,
            'text': text
        })

    # 提取表格
    tables = []
    for idx, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append({
            'index': idx,
            'data': table_data
        })

    # 合并所有文本
    all_text = '\n'.join([item['text'] for item in content])

    return {
        'sections': sections,
        'content': content,
        'tables': tables,
        'full_text': all_text
    }


def main():
    parser = argparse.ArgumentParser(description='读取Word需求文档')
    parser.add_argument('file_path', help='docx文件路径')
    parser.add_argument('-o', '--output', help='输出JSON文件路径')
    parser.add_argument('--text-only', action='store_true', help='只输出纯文本')

    args = parser.parse_args()

    if not Path(args.file_path).exists():
        print(f"Error: File not found: {args.file_path}")
        sys.exit(1)

    result = read_docx(args.file_path)

    if args.text_only:
        print(result['full_text'])
    elif args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        print(f"Output written to: {args.output}")
    else:
        print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == '__main__':
    main()
